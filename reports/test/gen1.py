from docx import Document 

document = Document() 
paragraph = document.add_paragraph("Lorem ipsum dolor sit amet.")
document.add_heading('The REAL meaning of the universe')
document.add_heading('The role of dolphins', level=2) 
document.add_page_break()

table = document.add_table(rows=2, cols=2) 
cell = table.cell(0,1)
cell.text = 'parrot, possibly dead'
row = table.rows[1] 
row.cells[0].text = 'Foo bar to you.' 
row.cells[1].text = 'And a hearty foo bar to you too sir!'
col = table.columns[1]
col.cells[1].text = "The light at the end of the tunnel"
# looping through the table rows
for row in table.rows: 
    for cell in row.cells: 
        print(cell.text) 
# count of rows
row_count = len(table.rows) 
col_count = len(table.columns)
# Add rows 
row = table.add_row()

# example 
items = (
    (7, '1024', 'Plush kittens'), 
    (3, '2042', 'Furbees'), 
    (1, '1288', 'French Poodle Collars, Deluxe'), 
)

table = document.add_table(1, 3)
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Qty'
heading_cells[1].text = 'SKU'
heading_cells[2].text = 'Description' 

for item in items: 
    cells = table.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = item[2]

table.style = 'LightShading-Accent1'

document.save('./doc1.docx')