from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Function to create a styled title
def add_title(doc, text):
    title = doc.add_paragraph()
    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_subtitle(doc, text):
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_paragraph(doc, text, font_size=11, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    paragraph.alignment = alignment

def add_table(doc, data, col_widths):
    table = doc.add_table(rows=1, cols=len(col_widths))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    # Add headers
    for i, header in enumerate(data[0]):
        hdr_cells[i].text = header

    # Add rows
    for row_data in data[1:]:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            row_cells[i].width = col_widths[i]

def create_report():
    doc = Document()

    # Title and subtitle
    add_title(doc, "Monthly Membership Report")
    add_subtitle(doc, "Prepared for: Praesidium Members")

    # Introduction
    add_paragraph(doc, "Date: December 2024", bold=True)
    add_paragraph(doc, "Summary of membership activities and financials for the month.")

    # Membership section
    add_paragraph(doc, "Membership", bold=True, font_size=13)
    add_paragraph(doc, "This section includes details on active members, new enrollments, and participation rates.")

    # Example table data
    membership_data = [
        ["Month", "Active Members", "New Members", "Participation Rate"],
        ["December", 120, 15, "85%"],
        ["November", 115, 10, "80%"],
    ]

    add_table(doc, membership_data, [1500, 2000, 2000, 2000])

    # Finances section
    add_paragraph(doc, "Finances", bold=True, font_size=13)
    add_paragraph(doc, "This section summarizes revenue and expenditures.")

    finances_data = [
        ["Month", "Revenue", "Expenditures", "Net Balance"],
        ["December", "$5000", "$3000", "$2000"],
        ["November", "$4500", "$2500", "$2000"],
    ]

    add_table(doc, finances_data, [1200, 1500, 1500, 1500])

    # Save the document
    doc.save("Monthly_Membership_Report.docx")

# Run the function to create the report
create_report()
