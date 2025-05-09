from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from math import ceil
from report_variables import *

doc = Document()

h = doc.add_heading(CURIA_NAME, level=2)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER 
h = doc.add_heading(CURIA_PARISH, level=2)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h = doc.add_heading("PRAESIDIUM REPORT", level=2)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")
p = doc.add_paragraph("E-MAIL: ")
p.add_run(CURIA_EMAIL).underline = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("DATE OF REPORT: ")
run.bold = True 
p.add_run(DATE_OF_REPORT)
p.add_run("\t\tDATE OF LAST REPORT: ").bold = True 
p.add_run(DATE_OF_LAST_REPORT)
# run.font.size = Pt(19)
# p.paragraph_format.alignment 

p = doc.add_paragraph("", style="List Number")
p.add_run("NAME OF PRAESIDIUM: ").bold = True 
p.add_run(PRAESIDIUM_NAME)

p = doc.add_paragraph("", style="List Number")
p.add_run("ADDRESS: ").bold = True 
p.add_run(PRAESIDIUM_ADDRESS)

p = doc.add_paragraph("", style="List Number")
p.add_run("REPORT NO: ").bold = True 
p.add_run(str(REPORT_NO))
p.add_run("\t\tDATE OF INAUGURATION: ").bold = True
p.add_run(DATE_OF_INAUGURATION)

p = doc.add_paragraph("", style="List Number")
p.add_run("DAY/TIME/VENUE OF MEETING: ").bold = True 
p.add_run(MEETING_INFO)

p = doc.add_paragraph("", style="List Number")
p.add_run("PERIOD OF REPORT: ").bold = True 
p.add_run(REPORT_PERIOD)

p = doc.add_paragraph("", style="List Number")
p.add_run("DETAILS OF LAST VISIT BY CURIA: ").bold = True 
p.add_run(LAST_CURIA_VISIT_INFO)

officers = [
    ("Office", "Name", "Elected", "Term"), 
    ("Spiritual Director", SPIRITUAL_DIRECTOR['name'].title(), SPIRITUAL_DIRECTOR['elected'], "-"), 
    ("President", PRESIDENT['name'].title(), PRESIDENT['elected'], PRESIDENT['term'].title()), 
    ('Vice President', VICE_PRESIDENT['name'].title(), VICE_PRESIDENT['elected'], VICE_PRESIDENT['term'].title()), 
    ('Secretary', SECRETARY['name'].title(), SECRETARY['elected'], SECRETARY['term'].title()),
    ('Treasurer', TREASURER['name'].title(), TREASURER['elected'], TREASURER['term'].title())
]
p = doc.add_paragraph("", style="List Number")
p.add_run("OFFICERS").bold = True 
table = doc.add_table(rows=6, cols=4)
for i, row in enumerate(table.rows): 
    for cell, item in zip(row.cells, officers[i]): 
        cell.text = item 

attendance_to_curia = [
    ("Office", "No. of Meetings Attended", "Total Meetings for the Period", "Current Year", "Previous Year"), 
    (
        'President',
        CURIA_ATTENDANCE['PRESIDENT']['current'], 
        CURIA_ATTENDANCE['PRESIDENT']['current_total'], 
        f"{CURIA_ATTENDANCE['PRESIDENT']['current']} out of {CURIA_ATTENDANCE['PRESIDENT']['current_total']}", 
        f"{CURIA_ATTENDANCE['PRESIDENT']['previous']} out of {CURIA_ATTENDANCE['PRESIDENT']['previous_total']}" if CURIA_ATTENDANCE['PRESIDENT']['previous'] else NON 
    ), 
    (
        'Vice President',
        CURIA_ATTENDANCE['VICE_PRESIDENT']['current'], 
        CURIA_ATTENDANCE['VICE_PRESIDENT']['current_total'], 
        f"{CURIA_ATTENDANCE['VICE_PRESIDENT']['current']} out of {CURIA_ATTENDANCE['VICE_PRESIDENT']['current_total']}", 
        f"{CURIA_ATTENDANCE['VICE_PRESIDENT']['previous']} out of {CURIA_ATTENDANCE['VICE_PRESIDENT']['previous_total']}" if CURIA_ATTENDANCE['VICE_PRESIDENT']['previous'] else NON
    ), 
    (
        'Secretary',
        CURIA_ATTENDANCE['SECRETARY']['current'], 
        CURIA_ATTENDANCE['SECRETARY']['current_total'], 
        f"{CURIA_ATTENDANCE['SECRETARY']['current']} out of {CURIA_ATTENDANCE['SECRETARY']['current_total']}", 
        f"{CURIA_ATTENDANCE['SECRETARY']['previous']} out of {CURIA_ATTENDANCE['SECRETARY']['previous_total']}" if CURIA_ATTENDANCE['SECRETARY']['previous'] else NON
    ), 
    (
        'Treasurer',
        CURIA_ATTENDANCE['TREASURER']['current'], 
        CURIA_ATTENDANCE['TREASURER']['current_total'], 
        f"{CURIA_ATTENDANCE['TREASURER']['current']} out of {CURIA_ATTENDANCE['TREASURER']['current_total']}", 
        f"{CURIA_ATTENDANCE['TREASURER']['previous']} out of {CURIA_ATTENDANCE['TREASURER']['previous_total']}" if CURIA_ATTENDANCE['TREASURER']['previous'] else NON
    )
]
doc.add_paragraph()
p = doc.add_paragraph("", style="List Number")
p.add_run("ATTENDANCE TO CURIA MEETINGS").bold = True 
table = doc.add_table(rows=5, cols=5)
for i, row in enumerate(table.rows): 
    for cell, item in zip(row.cells, attendance_to_curia[i]): 
        cell.text = str(item)  

attendance_to_praesidium = [
    ("Office", "No. of Meetings Attended", "Total Meetings for the Period", "Current Year", "Previous Year"), 
    (
        'President',
        PRAESIDIUM_ATTENDANCE['PRESIDENT']['current'], 
        PRAESIDIUM_ATTENDANCE['PRESIDENT']['current_total'], 
        f"{PRAESIDIUM_ATTENDANCE['PRESIDENT']['current']} out of {PRAESIDIUM_ATTENDANCE['PRESIDENT']['current_total']}", 
        f"{PRAESIDIUM_ATTENDANCE['PRESIDENT']['previous']} out of {PRAESIDIUM_ATTENDANCE['PRESIDENT']['previous_total']}" if PRAESIDIUM_ATTENDANCE['PRESIDENT']['previous'] else NON 
    ), 
    (
        'Vice President',
        PRAESIDIUM_ATTENDANCE['VICE_PRESIDENT']['current'], 
        PRAESIDIUM_ATTENDANCE['VICE_PRESIDENT']['current_total'], 
        f"{PRAESIDIUM_ATTENDANCE['VICE_PRESIDENT']['current']} out of {PRAESIDIUM_ATTENDANCE['VICE_PRESIDENT']['current_total']}", 
        f"{PRAESIDIUM_ATTENDANCE['VICE_PRESIDENT']['previous']} out of {PRAESIDIUM_ATTENDANCE['VICE_PRESIDENT']['previous_total']}" if PRAESIDIUM_ATTENDANCE['VICE_PRESIDENT']['previous'] else NON
    ), 
    (
        'Secretary',
        PRAESIDIUM_ATTENDANCE['SECRETARY']['current'], 
        PRAESIDIUM_ATTENDANCE['SECRETARY']['current_total'], 
        f"{PRAESIDIUM_ATTENDANCE['SECRETARY']['current']} out of {PRAESIDIUM_ATTENDANCE['SECRETARY']['current_total']}", 
        f"{PRAESIDIUM_ATTENDANCE['SECRETARY']['previous']} out of {PRAESIDIUM_ATTENDANCE['SECRETARY']['previous_total']}" if PRAESIDIUM_ATTENDANCE['SECRETARY']['previous'] else NON
    ), 
    (
        'Treasurer',
        PRAESIDIUM_ATTENDANCE['TREASURER']['current'], 
        PRAESIDIUM_ATTENDANCE['TREASURER']['current_total'], 
        f"{PRAESIDIUM_ATTENDANCE['TREASURER']['current']} out of {PRAESIDIUM_ATTENDANCE['TREASURER']['current_total']}", 
        f"{PRAESIDIUM_ATTENDANCE['TREASURER']['previous']} out of {PRAESIDIUM_ATTENDANCE['TREASURER']['previous_total']}" if PRAESIDIUM_ATTENDANCE['TREASURER']['previous'] else NON
    )
]
doc.add_paragraph()
p = doc.add_paragraph("", style="List Number")
p.add_run("ATTENDANCE TO PRAESIDIUM MEETINGS").bold = True 
table = doc.add_table(rows=5, cols=5)
for i, row in enumerate(table.rows): 
    for cell, item in zip(row.cells, attendance_to_praesidium[i]): 
        cell.text = str(item) 

# doc.add_page_break()

membership = [
    ("Membership", "Senior", "Junior"), 
    ("No. of Affiliated Praesidia (if any)", MEMBERSHIP['AFFILIATED_PRAESIDIA']['senior'], MEMBERSHIP['AFFILIATED_PRAESIDIA']['junior']), 
    ("Active Members", MEMBERSHIP['ACTIVE_MEMBERS']['senior'], MEMBERSHIP['ACTIVE_MEMBERS']['junior']), 
    ("Probationary Members", MEMBERSHIP['PROBATIONARY_MEMBERS']['senior'], MEMBERSHIP['PROBATIONARY_MEMBERS']['junior']), 
    ("Auxiliary Members", MEMBERSHIP['AUXILIARY_MEMBERS']['senior'], MEMBERSHIP['AUXILIARY_MEMBERS']['junior']), 
    ("Praetorian Members", MEMBERSHIP['PRAETORIAN_MEMBERS']['senior'], MEMBERSHIP['PRAETORIAN_MEMBERS']['junior']),
    ("Adjutorian Members", MEMBERSHIP['ADJUTORIAN_MEMBERS']['senior'], MEMBERSHIP['ADJUTORIAN_MEMBERS']['junior']),
    ("Total Members", MEMBERSHIP['TOTAL_MEMBERS']['senior'], MEMBERSHIP['TOTAL_MEMBERS']['junior']),
]
doc.add_paragraph('')
p = doc.add_paragraph("", style="List Number")
p.add_run("MEMBERSHIP").bold = True 
table = doc.add_table(rows=8, cols=3)
for i, row in enumerate(table.rows): 
    for cell, item in zip(row.cells, membership[i]): 
        cell.text = str(item) 


doc.add_paragraph('')
p = doc.add_paragraph("", style="List Number") 
p.add_run("MEETINGS AND ATTENDANCE").bold = True 
p = doc.add_paragraph("")
p.add_run("No. of meetings expected to be held: ").bold = True 
p.add_run(str(MEETINGS_ATTENDANCE['EXPECTED']))
p = doc.add_paragraph("")
p.add_run("No. of meetings held: ").bold = True 
p.add_run(str(MEETINGS_ATTENDANCE['HELD']))
p = doc.add_paragraph("")
p.add_run("Average attendance per meeting: ").bold = True 
p.add_run((str(MEETINGS_ATTENDANCE['AVG'])))
p = doc.add_paragraph("")
p.add_run("Reason for poor attendance: ").bold = True 
p.add_run((str(MEETINGS_ATTENDANCE['REASON'])))

# Active works 
doc.add_paragraph('')
p = doc.add_paragraph("", style="List Number") 
p.add_run("ACTIVE WORKS UNDERTAKEN").bold = True 
active_works = list(ACTIVE_WORKS.keys())
doc.add_paragraph(', '.join(active_works))

no_of_rows = ceil(len(active_works) / 2)
table = doc.add_table(rows=no_of_rows, cols=2)
work_count = 65
for i, work in enumerate(active_works):
    col = i % 2; row = i // 2
    cell = table.cell(row, col)
    details = ACTIVE_WORKS[work] 
    key_val = list(details.items())
    cell_text = f'{chr(work_count)}. {work}\n'
    for a, b in key_val: 
        cell_text += str(a) + ": " + str(b) + '\n'
    cell.text = cell_text 
    work_count += 1

# Other works 
doc.add_paragraph('')
p = doc.add_paragraph("", style="List Number") 
p.add_run("OTHER WORKS UNDERTAKEN").bold = True 
other_works = list(OTHER_WORKS.keys())
doc.add_paragraph(', '.join(other_works))
no_of_rows = ceil(len(other_works) / 2)
table = doc.add_table(rows=no_of_rows, cols=2)
work_count = 65
for i, work in enumerate(other_works):
    col = i % 2; row = i // 2
    cell = table.cell(row, col)
    details = OTHER_WORKS[work] 
    key_val = list(details.items())
    cell_text = f'{chr(work_count)}. {work}\n'
    for a, b in key_val: 
        cell_text += str(a) + ": " + str(b) + '\n'
    cell.text = cell_text 
    work_count += 1

# Achievements
doc.add_paragraph('')
p = doc.add_paragraph("", style="List Number") 
p.add_run("ACHIEVEMENTS RECORDED").bold = True 
num_rows = len(ACHIEVEMENTS)
table = doc.add_table(rows=num_rows, cols=3)
for i, key in enumerate(ACHIEVEMENTS.keys()):
    row = table.rows[i]
    row.cells[0].text = key
    row.cells[1].text = str(ACHIEVEMENTS[key]['current'])
    row.cells[2].text = str(ACHIEVEMENTS[key]['previous'])
    

# Plans for extension
doc.add_paragraph('')
p = doc.add_paragraph("", style="List Number") 
p.add_run("PLANS FOR EXTENSION WORK").bold = True 
doc.add_paragraph(EXTENSION_PLANS)

# Functions 
doc.add_paragraph('')
p = doc.add_paragraph('', style="List Number")
p.add_run('LEGION FUNCTIONS WITH ATTENDANCE').bold = True 
num_rows = len(FUNCTIONS)
table = doc.add_table(rows=num_rows, cols=4)
for i, key in enumerate(FUNCTIONS.keys()):
    row = table.rows[i]
    row.cells[0].text = key
    row.cells[1].text = str(FUNCTIONS[key]['date'])
    row.cells[2].text = str(FUNCTIONS[key]['current'])
    row.cells[3].text = str(FUNCTIONS[key]['previous'])


# functions = [
#     ('Function', 'Date', 'Current Year', 'Previous Year'), 
#     ('ACIES', '24/03/2024', 1, NON), 
#     ('MAY DEVOTION', '01/05/2024 - 31/05/2024', 4, NON), 
#     ('EDEL QUINN MASS', '12/05/2024', 17, NON), 
#     ('ANNUAL ENCLOSED RETREAT', '27/05/2024', 3, NON), 
#     ("MARY'S BIRTHDAY", '08/09/2024', 17, NON), 
#     ('OFFICERS WORKSHOP', '28/07/2024', 3, NON), 
#     ("OCTOBER DEVOTION", "01/10/2023 - 31/10/2023", 9, NON), 
#     ("DEPARTED LEGIONARIES MASS", "02/11/2023", 1, NON), 
#     ("FRANK DUFF MASS", "07/11/2023", 1, NON), 
#     ("LEGION CONGRESS", "12/06/2024", 2, NON), 
#     ("OUTDOOR FUNCTION", "28/08/2024", 5, NON), 
#     ("PATRICIANS MEETINGS", "04/08/2024", 7, NON), 
#     ("ANNUAL GENERAL REUNION (AGR)", "11/12/2023", 7, NON), 
#     ("EXPLORATIO DOMINICALIS", "29/06/2024", 2, NON) 
# ]
# # no_of_rows = ceil((len(functions)-1) / 2) # for later
# no_of_rows = len(functions)
# table = doc.add_table(rows=no_of_rows, cols=5)
# for i, row in enumerate(table.rows): 
#     cell = row.cells[0]
#     cell.text = str(i) 
#     for cell, item in zip(row.cells[1:], functions[i]): 
#         cell.text = str(item)  
# cell = table.cell(0, 0)
# cell.text = "S/N"

# Finance
doc.add_paragraph('')
p = doc.add_paragraph('', style="List Number")
p.add_run('FINANCE').bold = True 
num_rows = len(INCOME)+len(EXPENDITURE)+2
income_shift = len(INCOME) + 1
table = doc.add_table(rows=num_rows, cols=3)
row = table.rows[0] 
row.cells[0].text = 'DESCRIPTIONS'
row.cells[1].text = 'CREDIT'
row.cells[2].text = 'DEBIT'

for i in [j+1 for j in range(num_rows-1)]:   
    row = table.rows[i]    
    if i < income_shift: 
        # handle income 
        keys, vals = list(INCOME.keys()), list(INCOME.values())
        key, val = keys[i-1], vals[i-1]
        row.cells[0].text = str(key)
        row.cells[1].text = str(val)
    elif i < (num_rows-1): 
        # handle expenditure 
        keys, vals = list(EXPENDITURE.keys()), list(EXPENDITURE.values())
        key, val = keys[i-income_shift], vals[i-income_shift]
        row.cells[0].text = str(key)
        row.cells[2].text = str(val)
    else: 
        # handle last row 
        row.cells[0].text = "BALANCE AT HAND"
        row.cells[1].text = str(BALANCE) 

doc.add_paragraph('')
p = doc.add_paragraph('', style="List Number")
p.add_run('Was the finance audited? ').bold = True 
p.add_run(AUDITED_FINANCE_QUERY)

p = doc.add_paragraph('', style="List Number")
p.add_run('Problems: ').bold = True 
p.add_run(PROBLEMS)

p = doc.add_paragraph('', style="List Number")
p.add_run('Was the report read and accepted by members before submission? ').bold = True 
p.add_run(REPORT_ACCEPTED_BY_MEMBERS)

p = doc.add_paragraph('', style="List Number")
p.add_run('Any other comment/remark: ').bold = True 
p.add_run(COMMENTS)

p = doc.add_paragraph('', style="List Number")
p.add_run('Conclusion: ').bold = True 
p.add_run(CONCLUSION)

doc.add_paragraph()
p = doc.add_paragraph()
p1 = p.add_run('..................................................................')
p2 = p.add_run('\t\t\t..................................................................')

doc.add_paragraph()
p = doc.add_paragraph()
p1 = p.add_run(PRESIDENT['name'])
p2 = p.add_run('\t\t\t' + SECRETARY['name'])

doc.add_paragraph()
p = doc.add_paragraph()
p1 = p.add_run('President')
p2 = p.add_run('\t\t\tSecretary')

doc.add_paragraph("Spiritual Director's Comment:")
doc.add_paragraph('.'*180)
doc.add_paragraph('.......................')
doc.add_paragraph(SPIRITUAL_DIRECTOR['name'] + '\nSpiritual Director')


doc.save('test_rep2.docx')